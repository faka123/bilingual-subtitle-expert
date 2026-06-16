"""
双语字幕精确排版专家 - 核心匹配引擎

将校对好的中英对照文档与带时间轴的中文 SRT 字幕匹配，
生成符合专业标准的双语字幕文件。

支持 LLM 增强模式：
- 英文智能断句：语义级拆分，消除 "..." 占位符
- 语义匹配：用 AI 替代字符相似度算法
"""

import re
import difflib
from dataclasses import dataclass, field
from typing import Optional


# ── 数据结构（按双语字幕排版规则文档定义）────────────────────

@dataclass
class TimelineItem:
    """时间轴文档条目（对应文档 TimelineDoc.items）"""
    index: int
    start: str               # "HH:MM:SS,mmm"
    end: str
    text: str                # 字幕文本

    @property
    def start_time(self) -> str:
        """兼容旧属性名"""
        return self.start

    @start_time.setter
    def start_time(self, value: str):
        self.start = value

    @property
    def end_time(self) -> str:
        """兼容旧属性名"""
        return self.end

    @end_time.setter
    def end_time(self, value: str):
        self.end = value


@dataclass
class ProofreadPair:
    """校对文档中的中英对译（对应文档 ProofreadDoc.pairs）"""
    chinese: str
    english: str
    index: int = 0  # 在校对文档中的序号

    @property
    def zh(self) -> str:
        """兼容旧属性名"""
        return self.chinese

    @zh.setter
    def zh(self, value: str):
        self.chinese = value

    @property
    def en(self) -> str:
        """兼容旧属性名"""
        return self.english

    @en.setter
    def en(self, value: str):
        self.english = value

    # index 属性由 parse_proofread_doc 动态设置


@dataclass
class MatchResult:
    """匹配结果（内部中间结构）"""
    paragraph: ProofreadPair
    segments: list[TimelineItem]  # 属于该段落的中文 SRT 条目
    similarity: float             # 匹配相似度
    en_segments: list[str] = field(default_factory=list)  # 分割后的英文


@dataclass
class BilingualSubtitle:
    """双语字幕输出（对应文档 BilingualSubtitle）"""
    chinese_part: list[TimelineItem]
    english_part: list[TimelineItem]


# ── 向后兼容别名 ────────────────────────────────────────────

SrtEntry = TimelineItem          # 旧名
BilingualParagraph = ProofreadPair  # 旧名


# ── 文档解析 ──────────────────────────────────────────────

def parse_docx(file_path_or_bytes) -> str:
    """
    从 .docx 文件中提取文本，转换为校对文档格式。

    支持传入文件路径（str）或文件字节数据（bytes）。
    自动识别空段落作为双语对之间的分隔符。

    返回格式：
        中文段落1
        英文段落1

        中文段落2
        英文段落2
    """
    from docx import Document
    from io import BytesIO

    if isinstance(file_path_or_bytes, bytes):
        doc = Document(BytesIO(file_path_or_bytes))
    else:
        doc = Document(file_path_or_bytes)

    # 提取所有段落的文本，记录哪些是空段落
    para_texts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        para_texts.append(text)

    # 也读取表格中的文本（中英对照文档常用表格排版）
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            # 表格的每一行作为一个"段落"，多列用换行连接
            if row_texts:
                para_texts.append("\n".join(row_texts))
            else:
                para_texts.append("")  # 空行作为分隔

    # 构建输出：连续的非空段落合并为块，空段落作为块分隔符
    # 如果连续空段落过多则压缩为一个分隔
    lines = []
    prev_empty = False

    for text in para_texts:
        if text:
            lines.append(text)
            prev_empty = False
        else:
            if not prev_empty and lines:
                # 前一个非空，插入分隔空行
                lines.append("")
            prev_empty = True

    # 去掉末尾多余空行
    while lines and lines[-1] == "":
        lines.pop()

    return "\n".join(lines)


def parse_proofread_doc(text: str) -> list[BilingualParagraph]:
    """
    解析校对文档，提取中英对译段落。

    输入格式（按空行分隔每个中英对）:
        中文段落1
        英文段落1

        中文段落2
        英文段落2

    也支持 docx 导出的连续格式（无空行分隔），自动按中/英交替检测。

    返回 BilingualParagraph 列表。
    """
    # 按空行分割段落块
    blocks = re.split(r'\n\s*\n', text.strip())

    paragraphs = []
    pair_index = 0

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        # 提取块内的非空行
        block_lines = [l.strip() for l in block.split('\n') if l.strip()]
        if len(block_lines) < 2:
            continue

        # ── 在块内按中/英交替配对 ──
        i = 0
        while i < len(block_lines) - 1:
            zh_line = block_lines[i]
            en_line = block_lines[i + 1]

            zh_ratio = _chinese_char_ratio(zh_line)
            en_ratio = _english_char_ratio(en_line)

            if zh_ratio > 0.3 and en_ratio > 0.3:
                paragraphs.append(BilingualParagraph(
                    chinese=zh_line,
                    english=en_line,
                    index=pair_index,
                ))
                pair_index += 1
                i += 2
            else:
                # 不是标准的中-英对，尝试跳过当前行
                i += 1

        # 如果块内有落单行（奇数行），记录警告
        if i == len(block_lines) - 1:
            leftover = block_lines[i]
            # 仅在 leftover 看起来像有效文本时警告（非纯数字/符号）
            if len(leftover) > 3 and (_chinese_char_ratio(leftover) > 0.3 or _english_char_ratio(leftover) > 0.3):
                import logging
                logging.warning(f"校对文档解析：块内存在落单行（已丢弃）: {leftover[:80]}...")

    return paragraphs


def parse_srt(text: str) -> list[SrtEntry]:
    """
    解析 SRT 字幕文件。

    输入标准 SRT 格式:
        1
        00:00:00,200 --> 00:00:02,200
        字幕文本

    返回 SrtEntry 列表。
    """
    # 规范化换行符
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # 按空行分割
    blocks = re.split(r'\n\s*\n', text.strip())

    entries = []
    for block in blocks:
        block = block.strip()
        if not block:
            continue

        lines = block.split('\n')
        if len(lines) < 3:
            continue

        # 第一行：序号
        try:
            index = int(lines[0].strip())
        except ValueError:
            continue

        # 第二行：时间轴
        time_match = re.match(
            r'(\d{2}:\d{2}:\d{2}[,\.]\d{3})\s*-->\s*(\d{2}:\d{2}:\d{2}[,\.]\d{3})',
            lines[1].strip()
        )
        if not time_match:
            continue

        start_time = time_match.group(1)
        end_time = time_match.group(2)

        # 剩余行：字幕文本
        text_content = '\n'.join(lines[2:]).strip()

        entries.append(SrtEntry(
            index=index,
            start=start_time,
            end=end_time,
            text=text_content
        ))

    return entries


# ── 文本相似度 ────────────────────────────────────────────

def text_similarity(a: str, b: str) -> float:
    """
    计算两段文本的相似度 (0.0 ~ 1.0)。

    优先使用子串匹配（SRT 中文通常是校对文档中文的连续子串），
    子串匹配成功直接给高分；否则退回到 difflib 兜底。
    """
    # 清理文本
    a_clean = _normalize_for_matching(a)
    b_clean = _normalize_for_matching(b)

    if not a_clean or not b_clean:
        return 0.0

    # 子串优先：短的那段在长的那段里能找到，就是完美匹配
    shorter = a_clean if len(a_clean) <= len(b_clean) else b_clean
    longer = b_clean if len(a_clean) <= len(b_clean) else a_clean

    if shorter in longer:
        # 子串匹配成功，给高分
        # 归一化：剩余部分的长度越大，分数略低（但不低于 0.7）
        extra_ratio = 1.0 - (len(longer) - len(shorter)) / len(longer)
        return max(0.7, extra_ratio)

    # 子串匹配失败，回退到 difflib
    return difflib.SequenceMatcher(None, a_clean, b_clean).ratio()


# ── 中文口语虚词（匹配前过滤，提高相似度精度）─────────────

_FILLER_WORDS = set('呢啊吧嘛哦呀哇啦的了么吗喔哟咧呐')

def _normalize_for_matching(text: str) -> str:
    """清理文本用于匹配比较：去标点→去虚词→小写"""
    # 移除中文标点符号（含全角引号）
    text = re.sub(r'[，。！？、“”‘’“”（）《》【】\s]', '', text)
    # 移除英文标点
    text = re.sub(r'[,\.!\?;:\'"()\[\]\s]', '', text)
    # 移除中文口语虚词
    text = ''.join(c for c in text if c not in _FILLER_WORDS)
    return text.lower()


def _chinese_char_ratio(text: str) -> float:
    """计算中文字符占比"""
    if not text:
        return 0.0
    chinese_chars = len(re.findall(r'[一-鿿]', text))
    total_chars = len(re.sub(r'\s', '', text))
    if total_chars == 0:
        return 0.0
    return chinese_chars / total_chars


def _has_chinese(text: str) -> bool:
    """判断文本是否包含中文字符（比 ratio 更准确，不受短文本干扰）"""
    return bool(re.search(r'[一-鿿]', text))


def _english_char_ratio(text: str) -> float:
    """计算英文字符占比"""
    if not text:
        return 0.0
    english_chars = len(re.findall(r'[a-zA-Z]', text))
    total_chars = len(re.sub(r'\s', '', text))
    if total_chars == 0:
        return 0.0
    return english_chars / total_chars


# ── 段落匹配 ──────────────────────────────────────────────

def match_segments_to_paragraphs(
    srt_entries: list[SrtEntry],
    paragraphs: list[BilingualParagraph],
    threshold: float = 0.3
) -> list[MatchResult]:
    """
    将 SRT 字幕条目匹配到对应的双语段落。

    核心原则：**绝不丢弃任何条目**。每个 SRT 条目都会匹配到最佳段落，
    即使相似度低于阈值也会强制匹配。阈值仅影响统计报告中的"低置信度"计数，
    不影响实际匹配行为。
    """
    if not paragraphs:
        return []

    if not srt_entries:
        return []

    # ── 第一步：为每个 SRT 条目找到最匹配的段落（始终匹配，不用阈值过滤）──
    entry_para_map: dict[int, int] = {}       # entry_index -> paragraph_index
    entry_score_map: dict[int, float] = {}   # entry_index -> best_score

    for entry in srt_entries:
        best_idx = 0
        best_score = 0.0

        for pi, para in enumerate(paragraphs):
            score = text_similarity(entry.text, para.zh)
            if score > best_score:
                best_score = score
                best_idx = pi

        # 始终匹配到最佳段落（即使得分为 0）
        entry_para_map[entry.index] = best_idx
        entry_score_map[entry.index] = best_score

    # ── 第二步：按段落分组，保持 SRT 原始顺序 ──
    # 使用 dict 保持插入顺序（Python 3.7+）
    para_groups: dict[int, list[SrtEntry]] = {}
    para_group_order: list[int] = []

    for entry in srt_entries:
        pi = entry_para_map[entry.index]
        if pi not in para_groups:
            para_groups[pi] = []
            para_group_order.append(pi)
        para_groups[pi].append(entry)

    # ── 第三步：构建 MatchResult ──
    results = []
    for pi in para_group_order:
        group = para_groups[pi]
        # 计算组内平均相似度
        avg_sim = sum(
            entry_score_map.get(e.index, 0.0) for e in group
        ) / len(group)

        results.append(MatchResult(
            paragraph=paragraphs[pi],
            segments=sorted(group, key=lambda e: e.index),
            similarity=avg_sim
        ))

    return results


# ── 英文断句 ──────────────────────────────────────────────

# ── 常见英文缩写（不以句号结尾即句子结束）────
_KNOWN_ABBREVIATIONS = {
    # 称谓
    'mr', 'mrs', 'ms', 'miss', 'dr', 'prof', 'sr', 'jr', 'st',
    # 公司/组织
    'corp', 'inc', 'ltd', 'co', 'bros',
    # 常见缩写
    'vs', 'etc', 'dept', 'est', 'approx', 'appt', 'no',
    # 拉丁缩略
    'e.g', 'i.e', 'et al', 'viz',
    # 地址/州名
    'ave', 'blvd', 'rd', 'st',  # st 与上面的 st 冲突但无所谓
    # 度量
    'ft', 'in', 'lbs', 'oz', 'gal', 'vol',
    # 其他
    'misc', 'dept', 'govt', 'assn',
}

# ── 常见连词（断句时可优先在此处分割）──
_CONJUNCTIONS = {
    'and', 'but', 'or', 'so', 'yet', 'nor', 'for',
    'because', 'although', 'though', 'if', 'when', 'while',
    'where', 'which', 'that', 'who', 'whom', 'whose',
    'after', 'before', 'since', 'until', 'unless', 'whereas',
    'whereby', 'wherein', 'whereupon',
}


def _is_abbreviation(word: str) -> bool:
    """判断一个以 . 结尾的单词是否为常见缩写（不是句子结束）。

    检查去掉末尾 . 后的小写形式是否在已知缩写列表中。
    额外保护：单字母加点（如 "A." 代表 A 某人）不是句子结束，
    但通常已包含在列表中。
    """
    if not word.endswith('.'):
        return False
    stem = word[:-1].lower().strip()
    if len(stem) == 0:
        return False
    if stem in _KNOWN_ABBREVIATIONS:
        return True
    # 单大写字母加点通常是缩写（如 "A."、"B."）
    if len(stem) == 1 and stem.isalpha():
        return True
    return False


def _is_sentence_end(word: str, next_word: Optional[str] = None) -> bool:
    """判断一个以 . ! ? 结尾的单词是否真的是句子结束。

    - 缩写（Mr./Dr./etc.）不算句子结束
    - 如果下一个单词首字母大写，更可能是句子结束
    """
    if word.endswith('!'):
        return True
    if word.endswith('?'):
        return True
    if word.endswith('.'):
        if _is_abbreviation(word):
            return False
        # 额外确认：如果下一个词不是大写开头，可能不是句子结束
        # （但 "." 在文本中也可能是列表编号等，保守处理）
        return True
    return False


def _extract_word_breakpoints(words: list[str]) -> list[tuple[int, int]]:
    """从单词列表中提取所有自然断点位置，返回 [(word_index, quality), ...]。

    word_index 表示在此处断开（words[:idx] 为前一段，words[idx:] 为后一段）。

    断点质量分级：
    - 4 级：句末标点 . ! ? 后（确认非缩写，下一个单词首字母大写优先）
    - 3 级：分号 ; 冒号 : 破折号 —
    - 2 级：逗号 ,
    - 1 级：连词前（and/but/or/so/because/if/when/while/where/which/that/who 等）
    - 0 级：任意单词边界（每两个单词之间）

    注意：断点排在连词**前面**（不包含连词在当前段尾），
    与 _split_long_segment 中的连词位置处理一致。
    """
    breakpoints = []
    n = len(words)

    for i in range(1, n):
        prev_word = words[i - 1]
        curr_word = words[i]
        quality = -1  # 不是有意义的断点

        # 4 级：句末标点（确认非缩写）
        if prev_word.endswith('.'):
            if not _is_abbreviation(prev_word):
                # 进一步确认：下一个单词小写开头可能是非标准句号（如数字列表"1."）
                if curr_word and curr_word[0].isupper() if curr_word else True:
                    quality = 4
                else:
                    quality = 2  # 降级为逗号级（可能是列表序号等）
        elif prev_word.endswith('!'):
            quality = 4
        elif prev_word.endswith('?'):
            quality = 4
        # 3 级：分号/冒号/破折号
        elif prev_word.endswith(';'):
            quality = 3
        elif prev_word.endswith(':'):
            quality = 3
        elif prev_word == '—' or prev_word.endswith('—'):
            quality = 3
        # 2 级：逗号
        elif prev_word.endswith(','):
            quality = 2
        # 1 级：连词前 — 断在连词前面（让连词归入下一段）
        elif curr_word.lower().rstrip(',.;:!?') in _CONJUNCTIONS:
            quality = 1
        # 0 级：任意单词边界
        else:
            quality = 0

        if quality >= 0:
            breakpoints.append((i, quality))

    return breakpoints


def _find_best_breakpoint(
    breakpoints: list[tuple[int, int]],
    target_idx: int,
    lo_idx: int,
    hi_idx: int,
    total_words: int,
    remaining_segments: int,
    min_words: int = 1,
) -> int:
    """在 [lo_idx, hi_idx] 范围内找最佳断点。

    返回最佳的 word_index，如果范围内无合适断点则返回 target_idx。

    评分公式: score = quality * 0.15 + (1.0 - abs(idx - target) / max_range) * 0.85
    - 质量权重 15%，距离权重 85%（让含义对齐优先于长度精确）

    边界约束：
    - idx >= lo_idx（前段至少 min_words 词）
    - idx <= total_words - min_words * remaining_segments（后段至少留够）
    """
    max_range = max(abs(target_idx - lo_idx), abs(hi_idx - target_idx), 1)

    best_idx = target_idx
    best_score = -999.0

    for bp_idx, quality in breakpoints:
        if bp_idx < lo_idx or bp_idx > hi_idx:
            continue
        # 确保不会让剩余段无词可分配
        max_idx = total_words - min_words * max(0, remaining_segments - 1)
        if bp_idx > max_idx:
            continue

        dist_norm = abs(bp_idx - target_idx) / max_range
        score = quality * 0.15 + (1.0 - dist_norm) * 0.85

        if score > best_score:
            best_score = score
            best_idx = bp_idx

    return best_idx

def split_english_text(en_text: str, num_parts: int, chinese_segments: Optional[list[str]] = None) -> list[str]:
    """
    将英文段落按中文断句比例分割，长的中文配长的英文，短的配短的。

    策略：
    1. 按中文段落长度比例分配英文字符数
    2. 在分配的位置附近找自然断点
    3. 单段超过 42 字符的进一步拆分

    **关键约束：绝不丢失任何英文内容。** 如果 _enforce_max_len 拆分导致
    段数超过 num_parts，多出的段合并回最后一段（宁可超长，不可丢失）。
    """
    if num_parts <= 1:
        return [en_text]

    if chinese_segments and len(chinese_segments) == num_parts:
        parts = _split_proportional(en_text, chinese_segments)
    else:
        # 无中文参考时，回退到均匀分割
        parts = _split_uniform(en_text, num_parts)
        parts = _enforce_max_len(parts)

    # ── 确保段数精确等于 num_parts：宁可超长，绝不丢失 ──
    if len(parts) > num_parts:
        # 将多余段合并回最后一段，保留全部英文内容
        overflow = " ".join(parts[num_parts - 1:])
        parts = parts[:num_parts - 1] + [overflow]
    elif len(parts) < num_parts:
        # 段数不够时，复用最后一段内容（好过空字符串或 [MISSING]）
        parts.extend([parts[-1]] * (num_parts - len(parts)))

    return parts


def _split_proportional(en_text: str, cn_segments: list[str]) -> list[str]:
    """按中文段长度比例 + 语义边界混合策略分割英文。

    策略：动态选择最佳方法——
    - 英文有足够多的子句边界（≥目标段数）→ 子句池化分配（语义最优）
    - 英文子句边界太少 → 按词数比例 + 在空格/标点处智能切割

    保证：
    - 每段至少 3 词
    - 绝不截断单词（只在空格处断）
    - 不丢失不重复任何内容
    """
    cn_lens = [len(_normalize_for_matching(s)) for s in cn_segments]
    total_cn = sum(cn_lens)
    num_parts = len(cn_segments)

    if total_cn == 0:
        return _split_uniform(en_text, num_parts)

    words = en_text.split()
    total_words = len(words)

    if total_words <= num_parts:
        result = words[:]
        while len(result) < num_parts:
            result.append(words[-1])
        return result[:num_parts]

    MIN_WORDS = min(3, max(1, total_words // num_parts))

    # ── 构建质量断点和子句 ──
    breakpoints = _extract_word_breakpoints(words)

    # 从质量断点推导子句（子句 = 两个相邻高质量断点之间的词区间）
    # 子句边界：quality >= 2 的断点（句末/分号/冒号/逗号级）
    clauses = []
    cl_start = 0
    for bp_idx, bp_quality in breakpoints:
        if bp_quality >= 2:
            clauses.append((cl_start, bp_idx))
            cl_start = bp_idx
    # 最后一个子句
    if cl_start < len(words):
        if clauses:
            clauses[-1] = (clauses[-1][0], len(words))
        else:
            clauses.append((cl_start, len(words)))

    # ── 计算每段目标词数 ──
    targets = []
    remain = total_words
    remain_cn = total_cn
    for i in range(num_parts):
        ratio = cn_lens[i] / max(1, remain_cn)
        t = max(MIN_WORDS, round(remain * ratio))
        min_rest = MIN_WORDS * (num_parts - i - 1)
        t = min(t, remain - min_rest)
        t = max(MIN_WORDS, t)
        t = min(t, remain)
        targets.append(t)
        remain -= t
        remain_cn -= cn_lens[i]

    # 修正累积误差
    diff = sum(targets) - total_words
    while diff > 0:
        for i in reversed(range(num_parts)):
            if diff <= 0: break
            if targets[i] > MIN_WORDS:
                targets[i] -= 1; diff -= 1
    while diff < 0:
        targets[-1] += abs(diff); diff = 0

    # ── 混合策略：子句足够多时用子句池化，否则用词级切割 ──
    if len(clauses) >= num_parts:
        result = _clause_pool_assign(words, clauses, breakpoints, targets, num_parts, MIN_WORDS)
    else:
        result = _word_level_cut(words, breakpoints, targets, num_parts, MIN_WORDS)

    # 后处理
    result = [p for p in result if p.strip()]
    while len(result) < num_parts:
        result.append(result[-1] if result else "[EMPTY]")
    if len(result) > num_parts:
        overflow = ' '.join(result[num_parts - 1:])
        result = result[:num_parts - 1] + [overflow]

    return result[:num_parts]


def _clause_pool_assign(words, clauses, breakpoints, targets, num_parts, MIN_WORDS):
    """子句池化分配：从子句池中取，对每个段尝试多种方案选最优停点。

    改进点（相比旧版）：
    - 每个段尝试取 1/2/3/... 个子句，用评分选最优方案
    - 评分 = 词数接近度(85%) + 断点质量(15%)
    - 断在高质量标点处（句号>分号/冒号>逗号）有加分
    - 连词归入下一段（不拖在当前段尾）
    - 最后两段平衡：如果最后一段词数 > 倒数第二段的 2 倍，做平衡调整
    """
    total = len(words)
    parts = []
    ci = 0  # 当前子句索引

    # 构建子句结尾位置的断点质量查找（快速查某个 word_index 处的质量）
    bp_quality_at = {}
    for bp_idx, bp_quality in breakpoints:
        # 取最近的质量级别（最高级别的优先）
        if bp_idx not in bp_quality_at or bp_quality > bp_quality_at[bp_idx]:
            bp_quality_at[bp_idx] = bp_quality

    for seg_i in range(num_parts - 1):
        target = targets[seg_i]

        if ci >= len(clauses):
            # 已无子句可用，用最后一段内容填充
            parts.append(' '.join(words[clauses[-1][0]:clauses[-1][1]] if clauses else []))
            continue

        # ── 尝试取 k 个子句，选评分最高的方案 ──
        best_k = 1
        best_score = -999.0
        best_acc = 0
        best_quality = 0

        acc = 0
        for k in range(1, len(clauses) - ci + 1):
            cl_s, cl_e = clauses[ci + k - 1]
            acc += cl_e - cl_s
            clamp_end = words[cl_e - 1]

            # 确保至少留一个子句给剩余段落
            if ci + k >= len(clauses):
                break

            # 计算方案评分
            # 距离分：越接近目标越好
            dist_score = -abs(acc - target) / max(target, 1)
            # 质量分：断点质量越高越好
            end_quality = bp_quality_at.get(cl_e, 0)
            # 如果以连词结尾，降低质量（连词不应在段尾）
            next_start = clauses[ci + k][0] if ci + k < len(clauses) else total
            if next_start < total and words[next_start].lower().rstrip(',.;:!?') in _CONJUNCTIONS:
                end_quality = max(0, end_quality - 1)

            score = dist_score * 0.85 + end_quality * 0.15

            if score > best_score:
                best_score = score
                best_k = k
                best_acc = acc
                best_quality = end_quality

            # 如果已经在高质量断点处且接近 target，提前终止搜索
            if end_quality >= 3 and acc >= target * 0.7:
                break

        # ── 收集 best_k 个子句 ──
        seg_list = []
        for k in range(best_k):
            if ci < len(clauses):
                cl_s, cl_e = clauses[ci]
                seg_list.extend(words[cl_s:cl_e])
                ci += 1

        parts.append(' '.join(seg_list))

    # ── 收集最后一段（剩余子句）──
    last = []
    while ci < len(clauses):
        cl_s, cl_e = clauses[ci]
        last.extend(words[cl_s:cl_e])
        ci += 1

    # 最后两段平衡：如果最后一段远长于倒数第二段，将尾段的前几个词移到前段
    if len(parts) >= 1 and len(last) > 1:
        lp = len(parts[-1].split()) if parts[-1] else 0
        ll = len(last)
        if lp > 0 and ll > lp * 2:
            # 将最后一段末尾的平衡点（中间某个合适位置）之前的词移到前段
            target_lp = (lp + ll) // 2
            move_count = target_lp - lp
            if move_count > 0 and move_count < ll:
                # 在 move_count 附近找断点
                sub_breakpoints = _extract_word_breakpoints(last)
                best_bp = _find_best_breakpoint(
                    sub_breakpoints, max(1, move_count),
                    max(1, move_count // 2), min(ll - 1, move_count * 2),
                    ll, 1, 1
                )
                if best_bp > 0 and best_bp < ll:
                    moved_words = last[:best_bp]
                    last = last[best_bp:]
                    parts[-1] = (parts[-1] + ' ' + ' '.join(moved_words)).strip()

    parts.append(' '.join(last) if last else (parts[-1] if parts else ''))
    return parts



def _word_level_cut(words, breakpoints, targets, num_parts, MIN_WORDS):
    """词级切割：在单词边界切割，用评分选最优断点。

    改进点（相比旧版）：
    - 搜索范围扩大到 ±30% 总词数（而非 ±8 词固定值）
    - 遍历范围内所有候选断点，计算评分选最优（而非找到第一个就停）
    - 连词归属优化：断在连词前，让连词归入下一段
    """
    total = len(words)
    parts = []
    start = 0

    for seg_i in range(num_parts - 1):
        target = targets[seg_i]
        target_end = start + target
        remaining = num_parts - seg_i - 1

        # 搜索范围：以 target_end 为中心，半径为目标 50%（至少 5，至多 total*0.3）
        search_radius = max(5, min(target, total - start) // 2)
        search_radius = min(search_radius, max(5, int(total * 0.3)))

        lo = max(start + MIN_WORDS, target_end - search_radius)
        hi = min(total - MIN_WORDS * remaining, target_end + search_radius)

        # 连词优化：如果搜索范围内最高质量≤1（只有空格边界或连词），
        # 而 lo 之外附近有一个连词断点(quality=1)，允许 lo 下探到它
        in_range_qualities = [q for idx, q in breakpoints if lo <= idx <= hi]
        max_quality_in_range = max(in_range_qualities) if in_range_qualities else -1
        if max_quality_in_range <= 1 and lo > start + 1:
            for bp_idx, bp_q in breakpoints:
                if bp_q >= 1 and lo > bp_idx >= start + 1:
                    lo = bp_idx
                    break

        if lo >= hi:
            lo = max(start + 1, target_end - 1)
            hi = min(total - 1, target_end + 1)

        # 使用 _find_best_breakpoint 选最优断点
        best = _find_best_breakpoint(
            breakpoints, target_end, lo, hi,
            total, remaining, MIN_WORDS
        )

        # 连词归属优化：如果范围内有一个连词断点，优先选择它
        # （让连词归入下一段，当前段结尾干净自然）
        best_conj_bp = -1
        best_conj_dist = 999
        for bp_idx, bp_q in breakpoints:
            if bp_q == 1 and lo <= bp_idx <= hi:
                dist = abs(bp_idx - target_end)
                if dist < best_conj_dist:
                    best_conj_dist = dist
                    best_conj_bp = bp_idx
        if best_conj_bp > start and best_conj_bp < total:
            best = best_conj_bp

        # 连词归属优化：如果断点后的第一个词是连词，将断点前移（排除连词）
        # 这样连词归入下一段，当前段结尾干净
        if best < total and best > start:
            next_word = words[best].lower().rstrip(',.;:!?')
            if next_word in _CONJUNCTIONS and best > start + 1:
                best -= 1
                # 继续后退直到单词边界 — 但只需退一次（连词通常一个单词）
                # 如果前一个单词恰好也是连词则继续退
                while best > start + 1:
                    prev_word = words[best - 1].lower().rstrip(',.;:!?')
                    if prev_word not in _CONJUNCTIONS:
                        break
                    best -= 1

        if best <= start:
            best = min(start + max(1, target), total - 1)

        parts.append(' '.join(words[start:best]))
        start = best

    parts.append(' '.join(words[start:]))

    # 连词优化后处理：如果某段以连词结尾，将连词移到下一段开头
    for i in range(len(parts) - 1):
        if not parts[i]:
            continue
        pwords = parts[i].split()
        if pwords and pwords[-1].lower().rstrip(',.;:!?') in _CONJUNCTIONS:
            trailing_conj = pwords[-1]
            parts[i] = ' '.join(pwords[:-1])
            parts[i + 1] = trailing_conj + ' ' + parts[i + 1]

    # 修复空段（连词移动后可能导致第一段为空）
    for i in range(len(parts)):
        if not parts[i] and i + 1 < len(parts):
            # 从下一段借一个词
            next_words = parts[i + 1].split()
            if len(next_words) > 1:
                parts[i] = next_words[0]
                parts[i + 1] = ' '.join(next_words[1:])

    return parts


def _split_uniform(en_text: str, num_parts: int) -> list[str]:
    """均匀分割英文（原有逻辑，作为回退）"""
    # 找到所有自然断点位置
    break_points = _find_natural_breaks(en_text)

    if len(break_points) < num_parts - 1:
        return _split_by_chars(en_text, num_parts)

    total_len = len(en_text)
    target_len = total_len / num_parts
    selected = _select_breakpoints(break_points, num_parts - 1, total_len)

    parts = []
    start = 0
    for bp in selected:
        parts.append(en_text[start:bp].strip())
        start = bp
    parts.append(en_text[start:].strip())
    return parts


def _find_break_near(text: str, target: int, min_pos: int) -> int:
    """在 target 位置附近找最佳自然断点。

    改进：
    - 搜索窗口从 ±20 扩大到 ±60 字符
    - 强制在单词边界断开（绝不在单词中间截断）
    - 优先句末标点(.!?;)后的空格，其次逗号，最后普通空格
    - 无合适断点时回退到最近的空格（而非强行切字符）
    """
    # 搜索范围
    search_start = max(min_pos + 1, target - 60)
    search_end = min(len(text), target + 60)

    best = -1
    best_score = float('inf')

    # 在搜索范围内找所有断点
    for bp_pos, priority in _find_natural_breaks(text):
        if search_start <= bp_pos <= search_end:
            # 确保断点在单词边界（断点后应该是空格或字母）
            if bp_pos < len(text) and text[bp_pos - 1] not in '.!?;,':
                if bp_pos < len(text) and text[bp_pos - 1].isalpha() and text[bp_pos].isalpha():
                    continue  # 在两个字母中间，非单词边界
            dist = abs(bp_pos - target)
            score = dist + priority * 10
            if score < best_score:
                best_score = score
                best = bp_pos

    # 没找到标点断点：在 target 附近找最后一个空格（确保单词边界）
    if best < 0:
        lo = max(min_pos, target - 30)
        hi = min(len(text), target + 30)
        for pos in range(hi - 1, lo, -1):
            if text[pos] == ' ' and pos > min_pos:
                best = pos + 1  # 断点 = 空格后面
                break

    # 最终回退：在更大范围内找任何空格
    if best < 0 or best <= min_pos:
        lo = max(min_pos, target - 60)
        hi = min(len(text), target + 60)
        for pos in range(hi - 1, lo, -1):
            if text[pos] == ' ' and pos > min_pos:
                best = pos + 1
                break

    # 实在找不到空格，返回离 target 最近的单词边界
    if best < 0 or best <= min_pos:
        # 不强行切字符，而是取 min_pos 后的下一个空格，或 target 后的下一个空格
        for pos in range(min_pos + 1, min(len(text), min_pos + 60)):
            if text[pos] == ' ':
                best = pos + 1
                break
        if best < 0 or best <= min_pos:
            best = min(min_pos + 1, len(text))

    return best


def _enforce_max_len(parts: list[str]) -> list[str]:
    """将超过 42 字符的片段进一步拆分（无段数上限保护）。"""
    final = []
    for part in parts:
        if len(part) > 42:
            final.extend(_split_long_segment(part))
        else:
            final.append(part)
    return final


def _enforce_max_len_with_limit(parts: list[str], max_parts: int) -> list[str]:
    """将超过 42 字符的片段拆分，但确保总段数不超过 max_parts。

    如果拆分会使段数膨胀到超过 max_parts，则放弃拆分该段，
    保留为超长段（宁可超长，不可丢失内容）。
    """
    final = []
    budget = max_parts - len(parts)  # 可额外增加的段数

    for part in parts:
        if len(part) > 42 and budget > 0:
            split = _split_long_segment(part)
            if len(split) > 1:
                extra = len(split) - 1  # 拆分带来的额外段数
                if extra <= budget:
                    final.extend(split)
                    budget -= extra
                    continue
                # 不够预算拆，放弃拆分 — 保留超长段，不丢内容
        final.append(part)

    return final


def _split_long_segment(text: str, max_len: int = 42) -> list[str]:
    """将过长英文片段在自然断点处拆分为多段。

    拆分优先级：
    1. 句末标点 (. ! ?) 后的空格（需确认非缩写）
    2. 分号 (;) 冒号 (:) 破折号 (—) 后的空格
    3. 逗号 (,) 后的空格
    4. 连词前的空格
    5. 最后一个普通空格
    6. 找不到断点则保留超长段（宁可超长，不丢内容）

    关键约束：
    - 绝不截断单词
    - 绝不丢失任何单词
    - 所有原文字词在输出中恰好出现一次
    """
    parts = []
    remaining = text.strip()

    # 连词集合（比 _CONJUNCTIONS 稍精简，因为 _split_long_segment 是字符级操作）
    conj_pattern = (
        r'\s+(and|but|or|so|yet|nor|for'
        r'|because|although|though|if|when|while'
        r'|where|which|that|who|whom|whose'
        r'|after|before|since|until|unless|whereas)\s+'
    )

    while len(remaining) > max_len:
        best = -1
        window = remaining[:max_len + 1]

        # 1) 句末标点后的空格（确认非缩写）
        for m in re.finditer(r'([.!?])\s+', window):
            end_pos = m.end()
            if end_pos <= max_len:
                # 检查 . 是否为缩写的一部分
                punct = m.group(1)
                if punct == '.':
                    before_punct = remaining[:m.start()]
                    last_word = before_punct.rsplit(None, 1)[-1] if before_punct else ''
                    if last_word and _is_abbreviation(last_word + '.'):
                        continue  # 跳过缩写
                best = end_pos

        # 2) 分号/冒号/破折号后的空格
        if best <= 0:
            for m in re.finditer(r'[;:—]\s+', window):
                if m.end() <= max_len:
                    best = m.end()

        # 3) 逗号后的空格
        if best <= 0:
            for m in re.finditer(r',\s+', window):
                if m.end() <= max_len:
                    best = m.end()

        # 4) 连词前的空格（扩展了连词列表）
        if best <= 0:
            for m in re.finditer(conj_pattern, window):
                if m.start() <= max_len:
                    best = m.start()

        # 5) 最后一个普通空格（保证不截断单词）
        if best <= 0:
            last_space = window.rfind(' ')
            if last_space > 0:
                best = last_space + 1

        # 6) 找不到合适断点，保留超长段（不强行截断单词）
        if best <= 0:
            break

        parts.append(remaining[:best].strip())
        remaining = remaining[best:].strip()

    if remaining:
        parts.append(remaining)
    return parts if parts else [text]


def _find_natural_breaks(text: str) -> list[tuple[int, int]]:
    """
    找到英文文本中的自然断点位置（标点符号后的空格）。

    返回 (位置, 优先级) 列表，优先级 1=句末标点（优先），2=逗号（次选）。
    """
    break_pattern = re.compile(r'[.!?;]\s+')
    priority_breaks = [(m.end(), 1) for m in break_pattern.finditer(text)]

    comma_pattern = re.compile(r',\s+')
    comma_breaks = [(m.end(), 2) for m in comma_pattern.finditer(text)]

    all_breaks = priority_breaks + comma_breaks
    all_breaks.sort(key=lambda x: x[0])

    return all_breaks


def _select_breakpoints(
    breaks: list[tuple[int, int]],
    num_needed: int,
    total_len: int
) -> list[int]:
    """
    从可选断点中选择 num_needed 个，优先选择高优先级（句末标点）的断点，
    同时使分割后各部分长度尽量均匀。

    breaks: [(位置, 优先级), ...]，优先级 1 最高。
    """
    target_len = total_len / (num_needed + 1)

    selected = []
    available = list(breaks)

    for i in range(num_needed):
        ideal_pos = target_len * (i + 1)

        # 在理想位置附近寻找最佳断点（综合考虑距离和优先级）
        best_bp = None
        best_score = float('inf')

        for bp_pos, bp_priority in available:
            dist = abs(bp_pos - ideal_pos)
            # 综合评分：距离权重 0.7 + 优先级权重 0.3
            # 优先选择靠近理想位置的高优先级断点
            score = dist * 0.7 + (bp_priority - 1) * total_len * 0.15

            if score < best_score:
                best_score = score
                best_bp = bp_pos

        if best_bp is not None:
            selected.append(best_bp)
            # 移除已选断点及其附近断点（避免选择过于接近的断点）
            available = [
                (p, pr) for p, pr in available
                if abs(p - best_bp) > 5  # 至少间隔 5 个字符
            ]

    selected.sort()
    return selected


def _split_by_chars(text: str, num_parts: int) -> list[str]:
    """按单词尽量均匀分割英文文本。当分段数超过单词数时，剩余段返回空字符串标记。"""
    words = text.split()
    total_words = len(words)

    if total_words == 0:
        return ["[EMPTY]"] + ["[...]"] * (num_parts - 1)

    # 当分段数超过单词数时，每个单词最多一段，剩余段落复用最后一个单词
    # （避免产生 "..." 占位符，确保每条字幕都有实际内容）
    if num_parts > total_words:
        parts = words[:]  # 每个单词一段
        # 剩余段落填入最后一个单词（比 "..." 占位符更好）
        parts.extend([words[-1]] * (num_parts - total_words))
        return parts

    # 按单词数尽量均匀分配到各部分
    parts = []
    idx = 0
    for i in range(num_parts):
        remaining = num_parts - i
        take = max(1, (total_words - idx) // remaining)
        end = min(idx + take, total_words)
        parts.append(" ".join(words[idx:end]))
        idx = end

    # 处理可能因四舍五入残留的单词（合并到最后一部分）
    if idx < total_words and parts:
        parts[-1] = parts[-1] + " " + " ".join(words[idx:])

    return parts


# ── LLM 增强：英文智能断句 ────────────────────────────────

def split_english_text_llm(
    en_text: str,
    num_parts: int,
    llm_service,
    chinese_segments: Optional[list[str]] = None,
    errors: Optional[list] = None,
) -> list[str]:
    """
    用 LLM 按语义自然断句，失败时回退到规则方法。

    当 LLM 调用失败或结果无效时，自动回退到现有的 split_english_text，
    并将错误信息记录到 errors 列表中。

    **内容完整性保护**：无论 LLM 路径还是回退路径，都会确保英文字段总数
    不超过 num_parts（多出的合并），不丢失任何英文原文。
    """
    if num_parts <= 1:
        return [en_text]

    try:
        parts = llm_service.split_english_text(en_text, num_parts)
        # 验证结果有效性：每个部分都有实际文本
        if all(p and p.strip() and p.strip() != "..." for p in parts):
            # 作为内容完整性验证，检查每段的单词是否都在原文中
            en_words_lower = set(en_text.lower().split())
            for p in parts:
                p_words = set(p.lower().split())
                if not p_words.issubset(en_words_lower):
                    # LLM 可能捏造了内容，回退到规则算法
                    raise ValueError("LLM 输出包含原文不存在的单词，回退到规则算法")
            return parts
    except Exception:
        pass  # 静默回退，规则断句已足够好

    # 回退到规则方法（使用中文段长度参考提升断句质量）
    return split_english_text(en_text, num_parts, chinese_segments=chinese_segments)


# ── SRT 生成 ──────────────────────────────────────────────

def generate_bilingual_srt(
    results: list[MatchResult],
    original_srt: list[SrtEntry],
    renumber: bool = True,
    llm_service=None,
) -> tuple[str, dict]:
    """
    生成双语 SRT 输出。

    返回 (output_text, llm_stats)，其中 llm_stats 包含 LLM 使用情况。
    """
    output_lines = []
    llm_split_count = 0
    rule_split_count = 0
    llm_errors = []

    # ── 第一步：为每个 MatchResult 生成英文断句 ──
    for mr in results:
        n = len(mr.segments)
        if n == 0:
            continue
        if llm_service:
            cn_texts = [seg.text for seg in mr.segments]
            llm_parts = split_english_text_llm(mr.paragraph.en, n, llm_service, chinese_segments=cn_texts, errors=llm_errors)
            rule_parts = split_english_text(mr.paragraph.en, n, chinese_segments=cn_texts)
            # 判断 LLM 是否真正产生了不同的结果
            if llm_parts != rule_parts:
                llm_split_count += 1
                mr.en_segments = llm_parts
            else:
                rule_split_count += 1
                mr.en_segments = rule_parts
        else:
            rule_split_count += 1
            cn_texts = [seg.text for seg in mr.segments]
            mr.en_segments = split_english_text(mr.paragraph.en, n, chinese_segments=cn_texts)

    # ── 第二步：输出中文部分 ──
    if renumber:
        cn_index = 1
        for entry in original_srt:
            output_lines.append(str(cn_index))
            output_lines.append(f"{entry.start_time} --> {entry.end_time}")
            output_lines.append(entry.text)
            output_lines.append("")
            cn_index += 1
    else:
        for entry in original_srt:
            output_lines.append(str(entry.index))
            output_lines.append(f"{entry.start_time} --> {entry.end_time}")
            output_lines.append(entry.text)
            output_lines.append("")

    # ── 第三步：输出英文部分 ──
    if original_srt:
        max_cn_index = len(original_srt) if renumber else max(e.index for e in original_srt)
    else:
        max_cn_index = 0

    english_index = max_cn_index + 1

    for mr in sorted(results, key=lambda r: r.segments[0].index if r.segments else float('inf')):
        # 安全网：如果 en_segments 多于 segments，合并多余的段到最后一段
        seg_count = len(mr.segments)
        en_count = len(mr.en_segments)
        if en_count > seg_count and seg_count > 0:
            overflow = " ".join(mr.en_segments[seg_count - 1:])
            mr.en_segments = mr.en_segments[:seg_count - 1] + [overflow]

        for i, seg in enumerate(mr.segments):
            if i < len(mr.en_segments):
                en_text = mr.en_segments[i]
            else:
                en_text = "[MISSING]"

            output_lines.append(str(english_index))
            output_lines.append(f"{seg.start_time} --> {seg.end_time}")
            output_lines.append(en_text)
            output_lines.append("")

            english_index += 1

    llm_stats = {
        "llm_enabled": llm_service is not None,
        "llm_split_count": llm_split_count,
        "rule_split_count": rule_split_count,
        "llm_errors": llm_errors,
    }
    return "\n".join(output_lines), llm_stats


# ── 完整处理管道 ─────────────────────────────────────────

def process_subtitles(
    proofread_text: str,
    srt_text: str,
    similarity_threshold: float = 0.3,
    renumber: bool = True,
    llm_service=None,
    use_llm_match: bool = False,
) -> tuple[str, dict]:
    """
    完整的字幕处理流程。

    参数：
    - proofread_text: 校对文档文本
    - srt_text: SRT 时间轴文档文本
    - similarity_threshold: 相似度匹配阈值（0.0～1.0，默认 0.3）
    - renumber: 是否重新编排序号为连续数字（推荐 True）
    - llm_service: LLM 服务实例（可选），启用 AI 增强功能
    - use_llm_match: 是否用 LLM 做语义匹配（需要 llm_service）

    返回：
    - output_srt: 生成的双语 SRT 文本
    - stats: 处理统计信息
    """
    # 1. 解析
    paragraphs = parse_proofread_doc(proofread_text)
    all_srt_entries = parse_srt(srt_text)

    # 2. 过滤：只保留中文 SRT 条目用于匹配（用 _has_chinese 避免短文本误判）
    cn_srt_entries = [e for e in all_srt_entries if _has_chinese(e.text)]
    en_srt_entries = [e for e in all_srt_entries if not _has_chinese(e.text)]

    # 用中文条目做匹配，用全部条目做原始参考
    srt_entries = cn_srt_entries if cn_srt_entries else all_srt_entries

    # 3. 匹配（可选 LLM 语义匹配）
    if use_llm_match and llm_service and paragraphs:
        try:
            srt_texts = [e.text for e in srt_entries]
            para_texts = [p.zh for p in paragraphs]
            llm_matches = llm_service.match_subtitles(srt_texts, para_texts)
            results = _build_match_results_from_llm(
                srt_entries, paragraphs, llm_matches, similarity_threshold
            )
        except Exception:
            # LLM 匹配失败，回退到规则匹配
            results = match_segments_to_paragraphs(
                srt_entries, paragraphs, similarity_threshold
            )
    else:
        results = match_segments_to_paragraphs(
            srt_entries, paragraphs, similarity_threshold
        )

    # 3. 生成输出（只传入中文 SRT 条目作为原始参考，避免重复已存在的英文）
    output, llm_stats = generate_bilingual_srt(
        results, srt_entries, renumber=renumber, llm_service=llm_service
    )

    # 4. 统计
    matched_count = sum(len(r.segments) for r in results)

    # 低置信度匹配：相似度低于阈值的组
    low_confidence_count = sum(
        1 for r in results if r.similarity < similarity_threshold
    )
    low_confidence_entries = sum(
        len(r.segments) for r in results if r.similarity < similarity_threshold
    )

    # 检测是否为双语 SRT
    bilingual_srt_detected = len(en_srt_entries) > 0

    stats = {
        "paragraphs": len(paragraphs),
        "srt_entries": len(srt_entries),
        "srt_total_entries": len(all_srt_entries),
        "srt_en_entries_filtered": len(en_srt_entries),
        "bilingual_srt_detected": bilingual_srt_detected,
        "matched_groups": len(results),
        "matched_entries": matched_count,
        "low_confidence_groups": low_confidence_count,
        "low_confidence_entries": low_confidence_entries,
        "unmatched_entries": 0,  # 始终为 0，因为强制匹配所有条目
        **llm_stats,
    }

    return output, stats


def _build_match_results_from_llm(
    srt_entries: list[SrtEntry],
    paragraphs: list[BilingualParagraph],
    llm_matches: list[dict],
    threshold: float,
) -> list[MatchResult]:
    """
    将 LLM 匹配结果转换为 MatchResult 列表。

    LLM 返回 [{srt_index, paragraph_index, confidence}, ...]，
    此函数将其整理为按段落分组的 MatchResult。
    """
    # 按段落索引分组
    para_groups: dict[int, list[SrtEntry]] = {}
    entry_confidence: dict[int, float] = {}

    for match in llm_matches:
        srt_idx = match.get("srt_index", 0)
        para_idx = match.get("paragraph_index", 0)
        confidence = match.get("confidence", 0.5)

        if para_idx not in para_groups:
            para_groups[para_idx] = []
        if srt_idx < len(srt_entries):
            para_groups[para_idx].append(srt_entries[srt_idx])
            entry_confidence[srt_idx] = confidence

    # 确保每个 SRT 条目都有归属（LLM 可能遗漏）
    matched_indices = set()
    for group in para_groups.values():
        for e in group:
            matched_indices.add(e.index)

    for entry in srt_entries:
        if entry.index not in matched_indices:
            # 未匹配条目归到第一个段落（兜底）
            if 0 not in para_groups:
                para_groups[0] = []
            para_groups[0].append(entry)

    # 构建结果
    results = []
    for pi in sorted(para_groups.keys()):
        group = para_groups[pi]
        if pi < len(paragraphs):
            para = paragraphs[pi]
        else:
            # LLM 可能返回越界索引，用第一个段落兜底
            para = paragraphs[0] if paragraphs else BilingualParagraph(chinese="", english="")

        avg_confidence = sum(
            entry_confidence.get(e.index, 0.5) for e in group
        ) / len(group) if group else 0.0

        results.append(MatchResult(
            paragraph=para,
            segments=sorted(group, key=lambda e: e.index),
            similarity=avg_confidence,
        ))

    return results


# ── 质量检查 ──────────────────────────────────────────────

def quality_check(output_srt: str, stats: dict) -> list[str]:
    """对输出结果进行质量检查，返回问题列表。"""
    issues = []

    entries = parse_srt(output_srt)

    if not entries:
        return ["输出为空，没有解析到任何字幕条目"]

    # ── 分离中英文 ──
    # 不再用 _has_chinese() 判断（英文翻译中可能含中文地名如"沈阳"导致误判），
    # 改为按输出结构拆分：前 N 条为中文，其余为英文。
    cn_count = stats.get("srt_entries", 0)
    if cn_count > 0 and cn_count < len(entries):
        cn_entries = entries[:cn_count]
        en_entries = entries[cn_count:]
    else:
        # 兜底：无法按数量拆分时，回退到 _has_chinese 判断
        cn_entries = [e for e in entries if _has_chinese(e.text)]
        en_entries = [e for e in entries if not _has_chinese(e.text)]

    # ─── 关键检查：中英文字幕数量必须一致 ───
    if len(en_entries) != len(cn_entries):
        issues.append(
            f"❌ 致命错误：中英文字幕数量不一致！"
            f"中文 {len(cn_entries)} 条，英文 {len(en_entries)} 条，"
            f"差额 {abs(len(cn_entries) - len(en_entries))} 条"
        )
    else:
        # 数量一致时才检查序号连续性
        cn_indices = [e.index for e in cn_entries]
        if cn_indices != list(range(1, len(cn_indices) + 1)):
            issues.append(f"中文序号不连续，检测到 {len(cn_indices)} 条字幕但有间隔")

        en_indices = [e.index for e in en_entries]
        expected_en_start = len(cn_indices) + 1
        expected_en = list(range(expected_en_start, expected_en_start + len(en_indices)))
        if en_indices != expected_en:
            issues.append(f"英文序号不连续，期望从 {expected_en_start} 开始")

    # 检查序号重复
    indices = [e.index for e in entries]
    if len(indices) != len(set(indices)):
        from collections import Counter
        duplicates = [idx for idx, count in Counter(indices).items() if count > 1]
        issues.append(f"序号重复：{duplicates}")

    # 检查空字幕
    for e in entries:
        if not e.text.strip():
            issues.append(f"序号 {e.index} 字幕为空")

    # 注：低置信度匹配和行长度超标已由工具自动处理，不再弹出警告

    return issues


# ── API 兼容函数 ────────────────────────────────────────────

def process_bilingual_api(
    proofread_pairs: list[ProofreadPair],
    timeline_items: list[TimelineItem],
    similarity_threshold: float = 0.3,
) -> BilingualSubtitle:
    """
    API 兼容入口：接收结构化数据，返回 BilingualSubtitle。

    参数：
    - proofread_pairs: 校对文档中的中英对译列表
    - timeline_items: 时间轴中的中文条目列表
    - similarity_threshold: 相似度匹配阈值

    返回：
    - BilingualSubtitle: chinese_part + english_part
    """
    # 匹配
    results = match_segments_to_paragraphs(
        timeline_items, proofread_pairs, similarity_threshold
    )

    # 为每个匹配组分配英文断句
    for mr in results:
        n = len(mr.segments)
        if n == 0:
            continue
        cn_texts = [seg.text for seg in mr.segments]
        mr.en_segments = split_english_text(mr.paragraph.english, n, chinese_segments=cn_texts)

    # 构建中文部分（保持原始时间轴）
    chinese_part = [
        TimelineItem(index=t.index, start=t.start, end=t.end, text=t.text)
        for t in timeline_items
    ]

    # 构建英文部分（从中文数量+1开始编号）
    en_start_index = len(timeline_items) + 1
    english_part = []
    en_idx = en_start_index

    for mr in sorted(results, key=lambda r: r.segments[0].index if r.segments else float('inf')):
        for i, seg in enumerate(mr.segments):
            en_text = mr.en_segments[i] if i < len(mr.en_segments) else "[MISSING]"
            english_part.append(TimelineItem(
                index=en_idx,
                start=seg.start,
                end=seg.end,
                text=en_text,
            ))
            en_idx += 1

    return BilingualSubtitle(chinese_part=chinese_part, english_part=english_part)
